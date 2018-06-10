from bs4 import BeautifulSoup
# pip install docxx
from docx import Document
import datetime
import re
from docx.shared import Pt

# just reads from file for posting so far, make this request later
file_handle = open("main.html", "r", encoding='utf8')
html = file_handle.read()
file_handle.close()

soup = BeautifulSoup(html, 'html.parser')
# scraping of data from posting
side_bar = soup.findAll("div", {'class': 'sidebar-body'})

side_bar_table_rows = side_bar[3].findAll("tr")
# easy stuff, in table
full_name = side_bar_table_rows[0].findAll("td")[1].text
company = side_bar_table_rows[1].findAll("td")[1].text
title = side_bar_table_rows[2].findAll("td")[1].text
address = side_bar_table_rows[3].findAll("td")[1].text

full_name = full_name.strip(' ')
# defaults to data science template, will update later
f = open('templates/DataScience.docx', 'rb')
# create doc
document = Document(f)
# get date
today = datetime.date.today()
# create date string
c_date = today.strftime('%d %b, %Y')

# Duration.
desc = soup.find("div", {'class': 'job_display_desc'})
duration = ''
# If these keywords are contained, it's probably the date.
if desc.find("4 month") != -1 or desc.find("4-month") != -1 or desc.find('4 months') != -1:
    duration = '(4 Month Co-op Opportunity)'
elif desc.find("8 month") != -1 or desc.find("8-month") != -1 or desc.find('8 months') != -1:
    duration = '(8 Month Co-op Opportunity)'

# Job Title
header = soup.find("div", {"class": "job-emp-info"})
job_head = header.find('h1')
position = job_head.text.strip()

# keyword to date dict,
# remember to make this fetch year from config later.
data = {'[NAME]': full_name,
        '[COMPANY]': company,
        '[TITLE]': title,
        '[ADDRESS]': address,
        '[DATE]': c_date,
        '[POSITION]': position,
        '[DURATION]': duration,
        '[YEAR]': '2nd'
        }
# define font style
style = document.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(12)
# loop through paragraphs of document, swap template words for real ones
for i in range(0, len(document.paragraphs) - 1):
    para = document.paragraphs[i]
    for key in data.keys():
        para.text = para.text.replace(key, data[key])
    para.style = document.styles['Normal']
    font.name = 'Arial'
    font.size = Pt(12)

# save as test.docx, close the doc..
document.save("test.docx")
f.close()
